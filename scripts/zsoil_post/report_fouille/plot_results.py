# -*- coding: cp1252 -*-
# @description Plotting nodal results on cut section using vtk.
# @input zsoil results
# @output png
# @author Matthias Preisig
# @date 2018/06/22
import math
import numpy as np
import matplotlib.pyplot as plt
from matplotlib.patches import Polygon
from matplotlib.collections import PatchCollection
import vtk
import os,re
import pickle

from docx import Document
from docx.shared import Cm

import plot_disp_paroi
import plot_disp_evol
import plot_efforts_etais
import plot_efforts_evol_etais
import plot_efforts_longrines
import plot_cut_paroi
import plot_disp_cut_contours

from zsoil_tools import zsoil_results as zr
from zsoil_tools import vtktools
from zsoil_tools import postpro_lib as pl

pathname = '../'
pvpath = '../pv'
pblist = ['M1224_v4(secPiles)_3D_v1_1(ancN3)_(ULsubStep)_urgenceButon']

##############################################################################
# steps:
##############################################################################
class c_etapes:
    def __init__(self):
        self.names = [u'Exc. 705 msm',
                      u'Exc. 703 msm',
                      u'Exc. 698.5 msm',
                      u'Exc. 696.25 msm',
                      u'Exc. 694 msm',
                      u'Exc. 691.3 msm, clous sud',
                      u'F.f. atteint',
                      u'Suppression N3',
                      u'Suppression N2']
##                      u'Suppression N3','','','','','',
##                      u'Suppression N2','','','','','',
##                      u'Suppression N1']
        self.tvect = [3,5,7,8,9,10,11,13,14]#,15]
etapes = c_etapes()


##############################################################################
# coupes:
##############################################################################
class c_coupes:
    def __init__(self):
        self.names = ['Angle-M2','4-4','Etai 1 Ouest','Etai 1 Nord',
                      'B12 Nord','B12 Nord (sur étai)','Etai 3 Ouest']
        orig = [[26724.91,0,-11685.19],
                [26715.96,0,-11678.73],
                [26721.116,0,-11682.451],
                [26733.718,0,-11683.46],
                [26750.63,0,-11674.99],
                [26753.63,0,-11673.49],
                [26709.59,0,-11674.129]]
        self.origins = [[v[0]-26600,0,v[2]+11800] for v in orig]
        self.normals = [[0.82629588,0,-0.56323628],
                        [0.82629588,0,-0.56323628],
                        [0.82629588,0,-0.56323628],
                        [0.89442719,0,0.4472136],
                        [0.89442719,0,0.4472136],
                        [0.89442719,0,0.4472136],
                        [0.82629588,0,-0.56323628]]
coupes = c_coupes()

class c_coupes_vol:
    def __init__(self):
        self.names = [u'Paroi Ouest - entre étais 2 et 3',
                      u'Paroi Ouest - étais 3']
        self.origins = [[111.78,0,124.29],
                        [109.35,0,125.55]]
        self.normals = [[0.729702,0,-0.526817],
                        [0.729702,0,-0.526817]]
        self.bounds = [[[-20,5],
                        [690,708]],
                       [[-20,5],
                        [690,708]]]
coupes_vol = c_coupes_vol()

for kp,prob in enumerate(pblist):
    tx = []
    try:
        res = pickle.load(open(prob+'.p', "rb" ))
        
        print(prob+' loaded')
        tsteps = res.out_steps
        tvect = [res.steps[kt].time for kt in tsteps]
    except:
        res = zr(pathname,prob)
        res.read_rcf()
        res.read_his()
        tsteps = []
        for kt,step in enumerate(res.steps):
            if step.conv_status==-1:
                if step.time in etapes.tvect or abs(step.time%1)<1e-3:
                    tsteps.append(kt)
        res.out_steps = tsteps
        res.read_dat()
        res.read_s00()
        res.read_s03()
        res.read_s04()
        pickle.dump(res, open(prob+'.p', "wb" ) ,pickle.HIGHEST_PROTOCOL)
    try:
        f = open(pvpath+'/'+prob+'.pvd')
        kt = 0
        for line in f:
            if 'shell.vtu' in line:
##            if '_vo_' in line:
                t = float(line[line.find('timestep=')+9:].split('"')[1])
                if t in tvect:
                    tx.append(t)
        f.close()
        print(pvpath+'/'+prob+'.pvd loaded')
    except:
##        res = zr(pathname,prob)
##        res.read_rcf()
##        res.read_his()
##        tsteps = []
##        for kt,step in enumerate(res.steps):
##            if step.conv_status in [-1]:
##                if step.time in tvect:
##                    tsteps.append(kt)
##                    tx.append(step.time)
##        res.out_steps = tsteps
##        res.read_dat()
##        res.read_s00()
        for lab in res.ele_group_labels:
            if lab=='VOLUMICS':
                res.read_s01()  # volumics
            elif lab=='SHELLS':
                res.read_s02()  # volumics
        ptemp = os.getcwd()
        os.chdir(pvpath)
        pvdFile = vtktools.create_pvd(res)
        vtktools.write_vtu(res,shells=True,verbose=False,pvdFile=pvdFile)
        vtktools.write_vtu(res,vol=True,verbose=False,pvdFile=pvdFile)
        vtktools.save_pvd(pvdFile)
        os.chdir(ptemp)


    document = Document()

    document.add_heading('Report', 0)
    p = document.add_paragraph(u'Résultats pour le modèle %s'%(prob))

    document.add_heading(u'Déplacements', level=1)
    
    ketapes = [0,1,2,6,7,8]#,9]
    niveaux = [1000,700]
    figs = plot_disp_paroi.plot_disp_paroi(prob,coupes,etapes,ketapes,niveaux,res,figpath='disp_paroi')
    for fig in figs:
        document.add_picture(fig+'.png', width=Cm(15))

    ketapes = [6]#,9]
    figs = plot_disp_cut_contours.plot_disp_cut_contours(prob,coupes_vol,etapes,ketapes,pvpath,figpath='disp_contours')
    for fig in figs:
        document.add_picture(fig+'.png', width=Cm(15))

    fig = plot_disp_evol.plot_disp_evol(prob,etapes,res,figpath='disp_paroi_evol')
    document.add_picture(fig+'.png', width=Cm(15))

    document.add_heading(u'Efforts dans les étais', level=1)

    fig = plot_efforts_etais.plot_efforts_etais(prob,res,figpath='efforts_etais')
    document.add_picture(fig+'.png', width=Cm(15))
    fig = plot_efforts_evol_etais.plot_efforts_evol_etais(prob,res,figpath='efforts_etais_evol')
    document.add_picture(fig+'.png', width=Cm(15))

    document.add_heading(u'Efforts dans les longrines', level=1)

    figs = plot_efforts_longrines.plot_efforts_longrines(prob,res,figpath='efforts_longrines')
    document.add_picture(figs[0]+'.png', width=Cm(15))
    document.add_picture(figs[1]+'.png', width=Cm(15))
    document.add_picture(figs[2]+'.png', width=Cm(15))

    document.add_heading(u'Coupes', level=1)
    ketapes = [0,1,2,6]#,9]
    figs = plot_cut_paroi.plot_cut_paroi(prob,coupes,etapes,ketapes,pvpath,figpath='cut_paroi')
    for fig in figs:
        document.add_picture(fig+'.png', width=Cm(15))



##    document.add_paragraph('list item 1', style='List Bullet')
##    document.add_paragraph('list item 2', style='List Bullet')
##    document.add_paragraph('list item 3', style='List Bullet')
##
##    document.add_paragraph('enumerate typ 1-st item', style='List Number')
##    document.add_paragraph('enumerate typ 2-nd item', style='List Number')
##    document.add_picture('./plots_tutorial_1/MXX-AW(+).png', width=Inches(4.25))
##
##    document.add_page_break()

    document.save('report_%s.docx'%(prob))
        

            




